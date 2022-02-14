import os
import tkinter as tk
from tkinter import messagebox as mb
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# root config
root = tk.Tk()
root.title("Назначение невролога")
root.geometry("1280x640+40+40")
root.resizable(False, False)
maincolor = "#A8DE56"
root.config(bg=maincolor)

# dictionary
t1 = ["Тремора при осмотре нет. ", "Тремор рук при осмотре. ", "Тремор ног при осмотре. ",
      "Тремор подбородка при осмотре. ",
      "\nК осмотру относится адекватно. ",
      "\nК осмотру относится адекватно: усилением двигательной активности, концентрацией взора. ",
      "\nПри осмотре беспокойный. ", "\nК осмотру относится настороженно. ",
      "\nК осмотру относится негативно: плачет. ",
      "Быстро успокаивается на руках у матери. ", "Просьбы выполняет полностью. ", "Просьбы выполняет частично. ",
      "Просьбы не выполняет. ",
      "Ребёнок активный. ", "Рассматривает окружение. ", "Тянется к игрушке. ", "Коммуникабельный. ",
      "В глаза смотрит. ",
      "\nКожа желтушная. ", "\nКожа мраморная. ",
      "\nАстенического телосложения. ", "\nПониженного питания. ", "\nИзбыточного питания. ",
      "\nГолова долихоцефалической формы. ", "\nГолова округлой формы. ", "Голова уплощена в области затылка. ",
      "Голова уплощена в области темени. ",
      "Нахождение костей черепа. ", "Кости черепа плотные. ", "Голова с выраженными лобными буграми. ",
      "Голова с выраженным венозным рисунком. ",
      "\nГолова по средней линии. ", "\nГолова вправо. Уплотнения кивательных мышц нет. ",
      "\nГолова влево. Уплотнения кивательных мышц нет. ",
      "Короткая шея. ", "Поворот головы в обе стороны не ограничен. ",
      "\nАмниотические тяжи. ",
      "\nАнтимонголоидный разрез глаз. ", "\nПтоз. ", "\nЭкзофтальм. ", "\nЭпикант. ",
      "\nМакротия. ", "\nМикрогения. ", "\nПрогения. ", "\nПрогнатия. ",
      "\nВитилиго. ", "\nНевус. ",
      ]

var1 = []
for c in range(len(t1)):
    var1.append(tk.BooleanVar())

t2 = ["Движения глазных яблок координированы. Зрачки D=S. ",
      "Реакция зрачков на свет: прямая и содружественная/ живые. ", "Фотореакции сохранены. ",
      "Расходящееся косоглазие. ", "Сходящееся косоглазие. ",
      "\nФиксирует. ", "\nФиксирует кратковременно. ",
      "\nСледит хорошо, в полном объёме. ", "\nСледит. ", "\nСледит: отвлекается. ", "\nПрактически не следит. ",
      "\nНе следит. ",
      "Слабость конвергенции OD. ", "Слабость конвергенции OS. ", "Слабость конвергенции OI. ",
      "Слабость VI нерва справа. ", "Слабость VI нерва слева. ",
      "Установочных движений нет. ",
      "Грефе (+). ", "Грефе (-). ", "Нистагм в крайних отведениях. ", "Горизонтальный нистагм. ",
      "\nТочки выхода ветвей тройничного нерва безболезненные при пальпации. ",
      "\nРеакция на звук: сохранена. ", "\nРеакция на звук: комплекс оживления.", "\nРеакция на звук: поворот головы. ",
      "\nРеакция на звук: сомнительная.",
      "\nРеакция на звук: abs. ",
      "\nЛицо симметрично. ", "\nЛицо симметрично в покое и при плаче. ",
      "\nЛицо симметрично в покое и при мимических движениях. ", "\nАсимметрия лица. ",
      "Опущен угол нижней челюсти справа. ", "Опущен угол нижней челюсти слева. ", "Глазные щели: D>S. ",
      "Глазные щели: D<S. ",
      "Синева носогубного треугольника. ", "Слабость лицевой мускулатуры. ", "Рот приоткрыт. ",
      "Язык высунут. ", "Язык подвижен. ", "Язык по средней линии. ", "Язык отклонён вправо. ", "Язык отклонён влево. ",
      "Язык утолщен, напряжён. ", "Язык раздвоен на кончике. ", "Короткая уздечка языка. ",
      "Атрофии, фибриллярных и фасцикулярных подёргиваний нет. ",
      "Мягкое нёбо по средней линии. ", "Мягкое нёбо подвисает справа. ", "Мягкое нёбо подвисает слева. ",
      "\nГлоточный рефлекс (+). ",
      "\nГлотание и сосание не нарушено. ", "Голос звонкий. "
                                            "Пищевой интерес есть. ", "С ложки кормят. ", "Жуёт. ",
      "Неохотно жуёт твёрдую пищу. ", "Пьёт из чашки. ", "Пьёт из поильника. "
      ]

var2 = []
for c in range(len(t2)):
    var2.append(tk.BooleanVar())

t3 = ["Тонус мышц достаточный. ", "Тонус мышц снижен ",
      "в плечах, ", "в плечах: больше в правом, ", "в плечах: больше в левом, ",
      "в надплечьях, ", "в надплечьях: больше в правом, ", "в надплечьях: больше в левом, ",
      "в мышцах спины, ", "в мышцах живота, ", "в ягодицах, ", "в наружной поверхности стоп, ",
      "тонус мышц повышен ",
      "в конечностях, ", "в конечностях: больше в правых, ", "в конечностях: больше в левых, ",
      "в дистальных отделах, ",
      "в кистях рук, ", "в кистях рук: больше в правой, ", "в кистях рук: больше в левой, ",
      "в икроножных мышцах, ", "в икроножных мышцах: больше в правой ноге, ",
      "в икроножных мышцах: больше в левой ноге, ",

      "\nАсимметрия стояния плечевого пояса. ", "Правое плечо приподнято. ", "Левое плечо приподнято. ",
      "Болезненность надплечий. ",
      "\nКисти рук закрыты. ", "Разведение в ТБ суставах: полное. ", "Разведение в ТБ суставах: ограничено. ",
      "\nАтрофии нет. ", "Контрактур нет. ",
      "Контрактура коленных суставов. ", "Контрактура голеностопных суставов. ", "Контрактура локтевых суставов. ",
      "Рекурвация коленных суставов. ", "Рекурвация голеностопных суставов. ", "Рекурвация локтевых суставов. ",
      "Тугоподвижность в коленных суставах. ", "Тугоподвижность в голеностопных суставах. ",
      "Тугоподвижность в локтевых суставах. ",
      "Клонус стоп. ", "Гипотрофия мышц. "
      ]

var3 = []
for c in range(len(t3)):
    var3.append(tk.BooleanVar())

hightonus_index = t3.index("тонус мышц повышен ")
last_hightonus_index = t3.index("в икроножных мышцах: больше в левой ноге, ")

t4 = ["СР с рук: D=S живые. ", "СР с рук: D=S снижены. ", "СР с рук: D=S оживлены. ", "СР с рук: D=S спастичны. ",
      "СР с ног коленные: D=S живые. ", "СР с ног коленные: D=S снижены. ", "СР с ног коленные: D=S оживлены. ",
      "СР с ног коленные: D=S спастичны. ",
      "СР с ног ахилловы: D=S живые. ", "СР с ног ахилловы: D=S снижены. ", "СР с ног ахилловы: D=S оживлены. ",
      "СР с ног ахилловы: D=S спастичны. ",
      "СР с ног ахилловы: abs. ",
      "\nБрюшные: D=S живые. ", "Бабинский (+). ", "Бабинский (-). ", "Вендеровича (+). ", "Вендеровича (-). ",
      "Подошвенный (+). ", "Подошвенный (-). ", "Симптомы натяжения (-). ",
      "Сила мышц в руках: 3 балла. ", "Сила мышц в руках: 4 балла. ", "Сила мышц в руках: 5 баллов. ",
      "Сила мышц: не оценена из-за возраста ребёнка. ",
      "\nОпора на ноги (+): хорошая. ", "\nОпора на ноги: снижена. ", "\nОпора на ноги: abs. ",
      "База опоры расширена. ", "Ноги разогнуты в ТБ суставах. ", "Стопы вальгусные. ", "Поджимает пальцы стоп. ",
      "Cтопы ротированы внутрь. ", "\"Косолапит\" правой ногой. ", "\"Косолапит\" левой ногой. ",
      "Шаговый (+). ", "Шаговый: угасает. ", "Шаговый: попытка. ", "Шаговый: abs. ",
      "\"Заплетает\" правой стопой при шаговом. ", "\"Заплетает\" левой стопой при шаговом. ",
      "\nОпора на руки (+): хорошая. ", "\nОпора на руки: снижена. ", "\nОпора на руки: abs. ",
      "Опора на руки: на закрытых кистях. ",
      "Большие пальцы кистей открыты не полностью. ", "Первые и вторые пальцы кистей закрыты. ",
      "Ладушки играет. ", "Ладушки играет руками матери. ", "Ладушки не играет. ",
      "Игрушки берёт обеими руками. ", "Игрушки берёт активнее правой рукой. ", "Игрушки берёт активнее левой рукой. ",
      "Перекладывает игрушку из руки в руку. ", "Держит игрушку двумя руками. ",
      "Щипковый жест есть. ", "Пальцем показывает. ", "Захвата (+). ", "Захвата: снижен с правой руки. ",
      "Захвата: снижен с левой руки. ",
      ]

var4 = []
for c in range(len(t4)):
    var4.append(tk.BooleanVar())

t5 = ["При тракции за руки голову удерживает. ", "При тракции за руки голова провисает. ",
      "Пассивные перевороты не затруднены. ", "Переворачивается через оба бока. ",
      "Переворачивается через правый бок. ", "Переворачивается через левый бок. ",
      "Голову держит. ", "При вертикализации спину держит. ",
      "Сидит. ", "Сидит: спина круглая. ", "Сидит недолго, \"складывается\". ", "Не сидит. ", "Садится. ",
      "Самостоятельно не садится. ",
      "Ползает. ", "Ползает \"по-пластунски\". ", "Встаёт у опоры. ", "Ходит вдоль опоры. ",
      "Ходит с поддержкой за руки. ", "Ходит самостоятельно. ",
      "Походка уверенная. ",
      "Безусловные рефлексы: поисковый, сосательный, итд. ",
      "\nКоординаторные пробы не выполняет из-за возраста. ", "\nКоординаторные пробы выполняет уверенно. ",
      "В позе Ромберга устойчив. ",
      "Пальценосовую пробу выполняет уверенно. ", "Пальценосовую пробу выполняет с открытыми глазами. ",
      "\nЧувствительность не оценена из-за возраста ребёнка. ",
      "\nСоздаётся впечатление о сохранении чувствительности. ",
      "Ребёнок к опрятности приучен. ", "Навык опрятности: в стадии формирования. ",
      "Ребёнок к опрятности не приучен. ",
      "\nРечь: лепет. ", "\nРечь: гуление. ", "\nРечь: отдельные слоги. ", "\nРечь: отдельные слова. ",
      "\nРечь: фраза. ", "\nРечь: предложения. ",
      "Речь: чёткая. ", "Речь: с элементами дизартрии. ", "Речь: с элементами дислалии. ",
      "\nИгры: по-возрасту. ", "Разбрасывает игрушки. ", "Собирает игрушки. ", "Собирает пирамидку, пазлы. ",
      "Собирает конструктор \"LEGO\". ",
      "Катает машинку. ", "Качает куклу. ", "Имитирует звуки животных. "
      ]

var5 = []
for c in range(len(t5)):
    var5.append(tk.BooleanVar())

t6 = ["Здоров. ",
      "ЗЧМТ, СГМ. ",
      "Реконвалесцент ЗЧМТ, СГМ. ",
      "Данных за СГМ в настоящий момент нет. ",
      "Нечастые головные боли. ",
      "Нестабильность ШО. ",
      "Невротический синдром. ",
      "Астенический синдром. ",
      "ВПР. ",

      "ППЦНС. ",
      "ППЦНС: гипоксического генеза на фоне незрелости. ",
      "Асфиксия в родах. ",
      "Вентрикуломегалия. ",
      "Церебральная киста. ",
      "СЭК. ",
      "ППП. ",
      "ВЖК. ",
      "ВУИ. ",

      "ДЦП: спастический тетрапарез, диплегия. ",
      "ДЦП: спастический тетрапарез, гемиплегия. ",
      "Органическое поражение ЦНС. ",
      "Микроцефальный синдром. ",
      "Гипертензионный синдром. ",
      "Гидроцефальный синдром. ",
      "Гипертензионно-гидроцефальный синдром. ",
      "Реконвалесцент гипертензионного синдрома. ",
      "Реконвалесцент гидроцефального синдрома. ",
      "Реконвалесцент гипертензионно-гидроцефального синдрома. ",
      "Гидроцефалия. ",

      "Эпилепсия. ",
      "Эпиактивность по ЭЭГ. ",

      "Синдром мышечной дистонии: с кодом по МКБ. ",
      "Синдром мышечной дистонии: без кода по МКБ. ",
      "Синдром вялых надплечий. ",
      "Синдром пирамидной недостаточности в конечностях. ",
      "Нейрогенная кривошея. ",
      "Синдром двигательных нарушений. ",
      "Задержка моторного развития. ",

      "Гиперкинетический синдром. ",
      "Синдром гипервозбудимости. ",
      "Синдром срыгивания. ",
      "Фебрильные судороги. ",
      "Синдром вегетативных дисфункций. ",
      "ВСД. ",
      "Энурез. ",

      "Сходящееся косоглазие. ",
      "Расходящееся косоглазие. ",
      "Задержка психического развития. ",
      "Задержка речевого развития. ",
      "Задержка темпов речевого развития. ",
      "Заикание. ",
      "Дизартрия. ",
      "Диссомния. ",
      "Диссомния: прорезывание зубов. ",
      "Реакция на прорезывание зубов. ",
      "Раннее закрытие большого родничка. ",
      "Ранний восстановительный период. ",
      "Поздний восстановительный период. "
      ]

var6 = []
for c in range(len(t6)):
    var6.append(tk.BooleanVar())

t7 = ["Z 00.1: Рутинное обследование состояния здоровья ребенка.\n",
      "S 06.0: Сотрясение головного мозга.\n",
      "S 06.0: Сотрясение головного мозга.\n",
      "",
      "G 44.1: Сосудистая головная боль, не классифицированная в других рубриках.\n",
      "",
      "Z 73.3: Стрессовое состояние, не классифицированное в других рубриках.\n",
      "Z 73.3: Стрессовое состояние, не классифицированное в других рубриках.\n",
      "",

      "G 93.1: Аноксическое поражение головного мозга, не классифицированное в других рубриках.\n",
      "G 93.1: Аноксическое поражение головного мозга, не классифицированное в других рубриках.\n",
      "",
      "",
      "G 93.0: Церебральная киста.\n",
      "",
      "",
      "",
      "",

      "G 80.1: Спастический церебральный паралич, диплегия.\n",
      "G 80.2: Спастический церебральный паралич, гемиплегия.\n",
      "G 96.9: Поражение центральной нервной системы неуточнённое.\n",
      "",
      "G 93.2: Доброкачественная внутричерепная гипертензия.\n",
      "G 93.2: Доброкачественная внутричерепная гипертензия.\n",
      "G 93.2: Доброкачественная внутричерепная гипертензия.\n",
      "G 93.2: Доброкачественная внутричерепная гипертензия.\n",
      "G 93.2: Доброкачественная внутричерепная гипертензия.\n",
      "G 93.2: Доброкачественная внутричерепная гипертензия.\n",
      "G 91.8: Другие виды гидроцефалии.\n",

      "G 40.9: Эпилепсия неуточённая.\n",
      "",

      "G 24.8: Прочие дистонии.\n",
      "",
      "",
      "",
      "",
      "",
      "F 82: Специфические расстройства развития моторной функции.\n",

      "",
      "",
      "",
      "G 90.8: Другие расстройства вегетативной (автономной) нервной системы.\n",
      "G 90.8: Другие расстройства вегетативной (автономной) нервной системы.\n",
      "G 90.8: Другие расстройства вегетативной (автономной) нервной системы.\n",
      "F 98.0: Энурез неорганической природы.\n",

      "H 50.0: Сходящееся содружественное косоглазие.\n",
      "H 50.1: Расходящееся содружественное косоглазие.\n",
      "",
      "R 47.8: Другие и неуточнённые нарушения речи.\n",
      "R 47.8: Другие и неуточнённые нарушения речи.\n",
      "F 98.5: Заикание (запинание).\n",
      "R 47.1: Дизартрия и анартрия.\n",
      "G 47.9: Нарушение сна неуточнённое.\n",
      "K 00.7: Синдром прорезывания зубов.\n",
      "K 00.7: Синдром прорезывания зубов.\n",
      "",
      "",
      ""
      ]

t8 = ["дообследование в условиях 32 отделения ОДКБ им. Филатова\n",
      "дообследование в условиях 35 отделения ОДКБ им. Филатова\n",
      "дообследование в условиях дневного стационара ОДКБ им. Филатова\n",
      "консультация в Медико-генетическом научном центре\n",
      "консультации в НИИ педиатрии\n",

      "наблюдение педиатра\n", "наблюдение инфекциониста\n", "наблюдение окулиста\n", "наблюдение ортопеда\n",
      "наблюдение остеопата\n",
      "наблюдение эпилептолога\n", "наблюдение генетика\n", "наблюдение нейрохирурга\n", "наблюдение хирурга\n",
      "наблюдение кардиолога\n",
      "наблюдение гастроэнтеролога\n", "наблюдение нефролога\n", "наблюдение ЛОР-врача\n", "наблюдение сурдолога\n",
      "наблюдение психиатра\n",

      "консультация педиатра\n", "консультация инфекциониста\n", "консультация окулиста\n", "консультация ортопеда\n",
      "консультация остеопата\n",
      "консультация эпилептолога\n", "консультация генетика\n", "консультация нейрохирурга\n", "консультация хирурга\n",
      "консультация кардиолога\n"
      "консультация гастроэнтеролога\n", "консультация нефролога\n", "консультация ЛОР-врача\n",
      "консультация сурдолога\n", "консультация психиатра\n",

      "консультация нейропсихолога\n", "консультация онколога\n", "консультация сомнолога\n",

      "кровь на КФК, ЛДГ, АЛТ, АСТ, ЩФ\n", "кровь на ТТГ, Т-3, Т-4\n", "кровь на витамин D3\n",
      "кровь на кальций, кальций ионизированный\n",
      "НСГ (УЗИ головного мозга)\n", "ЭХО ЭС\n", "ЭЭГ: рутинная\n",
      "ЭЭГ мониторинг: 2 часа с включением дневного сна\n",
      "МРТ головного мозга\n", "УЗДГ сосудов шеи и головы\n", "ЭНМГ с верхних конечностей\n",
      "ЭНМГ с нижних конечностей\n",
      "рентгенография ШО\n", "прислать результат обследования на электронную почту \"Расту здоровым\""
      ]

var8 = []
for c in range(len(t8)):
    var8.append(tk.BooleanVar())

t9 = ["госпитализация в 35 отделение ОДКБ им. Филатова\n",
      "реабилитация в условиях стационара ОДКБ им. Филатова\n",
      "лечение в дневном стационаре ОДКБ им. Филатова\n",
      "охранительный режим\n",
      "ограничить: ТВ, компьютер, планшет, смартфон- до 10 минут в день\n",
      "режим сна и бодрствования\n",
      "формировать стереотип засыпания: закрывать шторы на время дневного сна, включать шумовой фон\n",
      "режим питания: каша за 1.5 часа до сна\n",
      "ЛФК на концентрацию взора и слуха\n",
      "лежать на животе 30-60-90 минут в день\n",
      "менять положение в пространстве\n",
      "снимать воротник Шанца\n",
      "упражнения на мяче: с фиксацией шейного отдела\n",
      "упражнения на мяче\n",
      "упражнения на мяче: лежать на мяче на животе, начинать \"подкатывать\"\n",
      "ЛФК на разведение конечностей, на раскрытие кистей: до кормления\n",
      "ЛФК на перевороты через оба бока\n",
      "ЛФК на удержание, перекладывание и перебирание игрушки\n",
      "ЛФК на укрепление плечевого пояса: стоять на руках, ходить на руках\n",
      "ЛФК на укрепление спины: носить с фиксацией за низ живота, затем нагибать\n",
      "ЛФК на укрепление ягодиц: носить с фиксацией за колени, затем нагибать\n",
      "ЛФК на навык ползания: выкладывать на пол с валиком под живот, учить раскачиваться\n",
      "жёсткая обувь: сидеть и ползать в обуви\n",
      "жёсткая обувь или босиком по неровной фактуре (ортопедические коврики, летом: песок, галька, трава)\n",
      "исключить ходунки\n",
      "подниматься по ступеням, пинать мяч, массажный мяч по наружной поверхности стоп\n",
      "ходить с отмашкой\n",
      "прыгать на двух ногах в два года, на одной ноге в три года\n",
      "занятия мелкой моторикой: перекладывать, перебирать, стучать, играть \"ладушки\"\n",
      "занятия мелкой моторикой: играть «ладушки», формировать щипковый и указательный жесты\n",
      "занятия мелкой моторикой: собирать, сортировать, складывать, лепить, рисовать, раскрашивать\n",
      "сочетанные движения руками\n",
      "пальчиковая гимнастика\n",
      "контрастные ванны для рук\n",
      "артикуляционная гимнастика: дуть, цокать, надувать щёки, имитировать животных, говорить коротко и чётко, жевать\n",
      "играть в ролевые игры\n",
      "формировать навык опрятности\n",
      "занятия с психологом\n",
      "занятия с логопедом\n",
      "занятия с дефектологом\n",
      "занятия с реабилитологом\n",
      "занятия по Войта, занятия по Бобат\n",
      "посещение бассейна\n",
      "посещение школ раннего развития\n",
      "ДДУ посещать может\n",
      "чередовать ванны: крапива, ромашка, череда\n",
      "успокаивающие ванны\n",
      "ванны с крапивой\n",
      "ванны с морской солью\n",
      "солёно-хвойные ванны\n",
      "витамин D3: по назначению педиатра\n",
      "витамин D3: не противопоказан со стороны нервной системы\n"
      ]

var9 = []
for c in range(len(t9)):
    var9.append(tk.BooleanVar())

t10 = ["массаж ", "по неврозу: воротниковая зона, волосистая часть головы, ",
       "по гипотонии: ",
       "плечи, ", "плечи (больше правое), ", "плечи (больше левое), ",
       "надплечья, ", "надплечья (больше правое), ", "надплечья (больше левое), ",
       "спина, ", "ягодицы, ", "наружная поверхность стоп, ", "лицо, ", "кисти рук, ",
       "по спастике: ",
       "конечности, ", "конечности (больше правые), ", "конечности (больше левые), ",
       "кисти рук, ", "кисти рук (больше правая), ", "в кисти рук (больше левая), ",
       "икроножные мышцы, ", "икроножные мышцы (больше правая), ", "икроножные мышцы (больше левая), ",
       "общий, ", "+ троксевазин гель в плечи при массаже\n",
       "парафин на кисти № 5+5\n",
       "парафиновые сапожки № 10\n",
       "парафин на щёки № 10\n",
       "электрофорез: 0.5% эуфиллин и 0.5% папаверин на шейный отдел и рукоятку грудины № 10\n",
       "электрофорез: 3% CaCl2 и 1% NaBr по Щербаку № 10\n",
       "электрофорез: 3% Cacl2 на коленные суставы № 10\n",
       "электрофорез: 0.1% атропин на область мочевого пузыря № 10\n",
       "гальванизация по Келлату № 10\n",
       "СМТ по гипотонии на межлопаточную область и правое надплечье № 7\n",
       "СМТ по гипотонии на межлопаточную область и левое надплечье № 7\n",
       "СМТ по гипотонии на межлопаточную область и надплечья № 7\n",
       "СМТ по гипотонии на длинные мышцы спины № 7\n",
       "СМТ по гипотонии на ягодицы № 7\n",
       "СМТ по гипотонии на наружный край голени и стопы обеих ног № 7\n",
       "СМТ по спастике на икроножные мышцы № 10\n",
       "СМТ по спастике на правую икроножную мышцу № 10\n",
       "CМТ по спастике на левую икроножную мышцу № 10\n",
       "СМТ по Квиташу № 10\n",
       "магнит транскраниально по гипертензионному синдрому № 10\n",
       "магнит на икроножные мышцы № 10\n",
       "лазер на шейный отдел № 5\n"
       ]

var10 = []
for c in range(len(t10)):
    var10.append(tk.BooleanVar())

common_index = t10.index("общий, ")

t11 = ["актовегин (Deproteinized hemoderivate de vitulos sanguinem) 200 мг",
       "актовегин (Deproteinized hemoderivate de vitulos sanguinem) внутримышечно",
       "атаракс (Hydroxyzine) 25 мг",
       "глиатилин (Choline alfoscerate) 600 мг/7 мл в питьевом растворе",
       "глицин (Glycine) 100 мг",
       "дантинорм бэби или дормикинд",
       "диакарб (Acetazolamide) 250 мг",
       "аспаркам (Potassium aspartate and magnesium aspartate) 175 мг",
       "кавинтон (Vinpocetine) 5 мг",
       "калгель или камистад бэби",
       "корилип-нео (Cocarboxylase + Riboflavin + Thioctic acid) свечи",
       "кортексин (Polypeptides de cerebri cortex donatus) 5 мг в 1.0 мл воды для инъекций",
       "кортексин (Polypeptides de cerebri cortex donatus) 10 мг в 2.0 мл воды для инъекций",
       "кудесан (Ubidecarenone) 30 мг/мл в питьевом растворе",
       "лист брусники",
       "магне В6 (Magnesii lactas + Pyridoxinum) в питьевом растворе",
       "мексиприм или мексидол (Ethyl-methyl-hydroxypyridine succinate) 125 мг",
       "мидокалм (Tolperisonum) 50 мг",
       "мочегонный чай",
       "мульти-табс бэби",
       "нервохель",
       "нооклерин (Deanol aceglumate) 200 мг/мл в питьевом растворе",
       "пантогам (Hopantenic acid) 100 мг/мл в питьевом растворе",
       "пантогам или пантокальцин (Hopantenic acid) 250 мг",
       "пикамилон (Nicotinoyl gamma-aminobutyric acid) 20 мг",
       "семакс (Methionyl-glutamyl-histidyl-phenylalanyl-prolyl-glycyl-proline) 0.1%",
       "спаскупрель",
       "танакан (Ginkgo Bilobae foliorum extract) 40 мг",
       "тенотен детский",
       "тералиджен (Alimemazine tartrate) 5 мг",
       "траумель",
       "фенибут (Aminophenylbutyric acid) 250 мг",
       "циннаризин или стугерон (Cinnarizine) 25 мг",
       "церебрум композитум",
       "цитофлавин (Inosine + Nicotinamide + Riboflavin + Succinic acid) внутривенно",
       "убихинон композитум",
       "эдас-306",
       "элькар (Levocarnitine) 300 мг/мл в питьевом растворе",
       "на ВКК, вакцинация"
       ]

var11 = []
for c in range(len(t11)):
    var11.append(tk.BooleanVar())

t12 = []


# first screen- patient info
def get_entry_start():
    global name, birthday, petition, head, chest, fonticulus, allergy
    name = ent_name.get().strip().title()
    birthday = ent_birthday.get().strip().replace(" ", ".")
    petition = ent_petition.get().strip().capitalize()
    head = ent_head.get().strip()
    chest = ent_chest.get().strip()
    fonticulus = ent_fonticulus.get().strip()
    allergy = ent_allergy.get().strip().capitalize()

    destroy_list = [lbl_name, lbl_birthday, lbl_petition, lbl_head, lbl_chest, lbl_fonticulus, lbl_allergy,
                    ent_name, ent_birthday, ent_petition, ent_head, ent_chest, ent_fonticulus, ent_allergy, btn]
    for object in destroy_list:
        object.destroy()
    checks_1()


# second screen- status
def checks_1():
    global checks1, btn
    checks1 = []
    for c in range(len(t1)):
        checks1.append(
            tk.Checkbutton(root, text=f"{t1[c][:60]}".replace("\n", ""), bg=maincolor, anchor="w", padx=10, width=52,
                           variable=var1[c],
                           onvalue=1, offvalue=0))
        if c < 20:
            rw = c;
            col = 0
        elif c >= 20 and c < 40:
            rw = c - 20;
            col = 1
        elif c >= 40 and c < 60:
            rw = c - 40;
            col = 2
        checks1[c].grid(row=rw, column=col)
    btn = tk.Button(root, text="Далее", bd=5, font=("Times New Roman", 12), width=10, height=2, command=get_entry_1)
    btn.grid(row=20, column=0, padx=20, pady=40)


def get_entry_1():
    for c in range(len(t1)):
        if var1[c].get() == False:
            t1[c] = ""
    for object in checks1:
        object.destroy()
    btn.destroy()
    checks_2()


def checks_2():
    global checks2, btn
    checks2 = []
    for c in range(len(t2)):
        checks2.append(
            tk.Checkbutton(root, text=f"{t2[c][:60]}".replace("\n", ""), bg=maincolor, anchor="w", padx=10, width=52,
                           variable=var2[c],
                           onvalue=1, offvalue=0))
        if c < 20:
            rw = c;
            col = 0
        elif c >= 20 and c < 40:
            rw = c - 20;
            col = 1
        elif c >= 40 and c < 60:
            rw = c - 40;
            col = 2
        checks2[c].grid(row=rw, column=col)
    btn = tk.Button(root, text="Далее", bd=5, font=("Times New Roman", 12), width=10, height=2, command=get_entry_2)
    btn.grid(row=20, column=0, padx=20, pady=40)


def get_entry_2():
    for c in range(len(t2)):
        if var2[c].get() == False:
            t2[c] = ""
    for object in checks2:
        object.destroy()
    btn.destroy()
    checks_3()


def checks_3():
    global checks3, btn
    checks3 = []
    for c in range(len(t3)):
        checks3.append(
            tk.Checkbutton(root, text=f"{t3[c][:60]}".replace("\n", ""), bg=maincolor, anchor="w", padx=10, width=52,
                           variable=var3[c],
                           onvalue=1, offvalue=0))
        if c < 20:
            rw = c;
            col = 0
        elif c >= 20 and c < 40:
            rw = c - 20;
            col = 1
        elif c >= 40 and c < 60:
            rw = c - 40;
            col = 2
        checks3[c].grid(row=rw, column=col)
    btn = tk.Button(root, text="Далее", bd=5, font=("Times New Roman", 12), width=10, height=2, command=get_entry_3)
    btn.grid(row=20, column=0, padx=20, pady=40)


def get_entry_3():
    for c in range(len(t3)):
        if var3[c].get() == False:
            t3[c] = ""
    for object in checks3:
        object.destroy()
    btn.destroy()
    checks_4()


def checks_4():
    global checks4, btn
    checks4 = []
    for c in range(len(t4)):
        checks4.append(
            tk.Checkbutton(root, text=f"{t4[c][:60]}".replace("\n", ""), bg=maincolor, anchor="w", padx=10, width=52,
                           variable=var4[c],
                           onvalue=1, offvalue=0))
        if c < 20:
            rw = c;
            col = 0
        elif c >= 20 and c < 40:
            rw = c - 20;
            col = 1
        elif c >= 40 and c < 60:
            rw = c - 40;
            col = 2
        checks4[c].grid(row=rw, column=col)
    btn = tk.Button(root, text="Далее", bd=5, font=("Times New Roman", 12), width=10, height=2, command=get_entry_4)
    btn.grid(row=20, column=0, padx=20, pady=40)


def get_entry_4():
    for c in range(len(t4)):
        if var4[c].get() == False:
            t4[c] = ""
    for object in checks4:
        object.destroy()
    btn.destroy()
    checks_5()


def checks_5():
    global checks5, btn
    checks5 = []
    for c in range(len(t5)):
        checks5.append(
            tk.Checkbutton(root, text=f"{t5[c][:60]}".replace("\n", ""), bg=maincolor, anchor="w", padx=10, width=52,
                           variable=var5[c],
                           onvalue=1, offvalue=0))
        if c < 20:
            rw = c;
            col = 0
        elif c >= 20 and c < 40:
            rw = c - 20;
            col = 1
        elif c >= 40 and c < 60:
            rw = c - 40;
            col = 2
        checks5[c].grid(row=rw, column=col)
    btn = tk.Button(root, text="Далее", bd=5, font=("Times New Roman", 12), width=10, height=2, command=get_entry_5)
    btn.grid(row=20, column=0, padx=20, pady=40)


def get_entry_5():
    for c in range(len(t5)):
        if var5[c].get() == False:
            t5[c] = ""
    for object in checks5:
        object.destroy()
    btn.destroy()
    diagnosis()


# third screen- diagnosis
def diagnosis():
    global checks6, btn
    checks6 = []
    for c in range(len(t6)):
        checks6.append(
            tk.Checkbutton(root, text=f"{t6[c][:60]}", bg=maincolor, anchor="w", padx=10, width=52, variable=var6[c],
                           onvalue=1, offvalue=0))
        if c < 20:
            rw = c;
            col = 0
        elif c >= 20 and c < 40:
            rw = c - 20;
            col = 1
        elif c >= 40 and c < 60:
            rw = c - 40;
            col = 2
        checks6[c].grid(row=rw, column=col)
    btn = tk.Button(root, text="Далее", bd=5, font=("Times New Roman", 12), width=10, height=2, command=get_entry_6)
    btn.grid(row=20, column=0, padx=20, pady=40)


def get_entry_6():
    for c in range(len(t6)):
        if var6[c].get() == False:
            t6[c] = ""
            t7[c] = ""
    for object in checks6:
        object.destroy()
    btn.destroy()
    prescription_1()


# fourth screen- prescription
def prescription_1():
    global checks8, btn
    checks8 = []
    for c in range(len(t8)):
        checks8.append(
            tk.Checkbutton(root, text=f"{t8[c][:60]}".replace("\n", ""), bg=maincolor, anchor="w", padx=10, width=52,
                           variable=var8[c],
                           onvalue=1, offvalue=0))
        if c < 20:
            rw = c;
            col = 0
        elif c >= 20 and c < 40:
            rw = c - 20;
            col = 1
        elif c >= 40 and c < 60:
            rw = c - 40;
            col = 2
        checks8[c].grid(row=rw, column=col)
    btn = tk.Button(root, text="Далее", bd=5, font=("Times New Roman", 12), width=10, height=2, command=get_entry_8)
    btn.grid(row=20, column=0, padx=20, pady=40)


def get_entry_8():
    for c in range(len(t8)):
        if var8[c].get() == False:
            t8[c] = ""
    for object in checks8:
        object.destroy()
    btn.destroy()
    prescription_2()


def prescription_2():
    global checks9, btn
    checks9 = []
    for c in range(len(t9)):
        checks9.append(
            tk.Checkbutton(root, text=f"{t9[c][:60]}".replace("\n", ""), bg=maincolor, anchor="w", padx=10, width=52,
                           variable=var9[c],
                           onvalue=1, offvalue=0))
        if c < 20:
            rw = c;
            col = 0
        elif c >= 20 and c < 40:
            rw = c - 20;
            col = 1
        elif c >= 40 and c < 60:
            rw = c - 40;
            col = 2
        checks9[c].grid(row=rw, column=col)
    btn = tk.Button(root, text="Далее", bd=5, font=("Times New Roman", 12), width=10, height=2, command=get_entry_9)
    btn.grid(row=20, column=0, padx=20, pady=40)


def get_entry_9():
    for c in range(len(t9)):
        if var9[c].get() == False:
            t9[c] = ""
    for object in checks9:
        object.destroy()
    btn.destroy()
    prescription_3()


def prescription_3():
    global checks10, btn
    checks10 = []
    for c in range(len(t10)):
        checks10.append(
            tk.Checkbutton(root, text=f"{t10[c][:60]}".replace("\n", ""), bg=maincolor, anchor="w", padx=10, width=52,
                           variable=var10[c],
                           onvalue=1, offvalue=0))
        if c < 20:
            rw = c;
            col = 0
        elif c >= 20 and c < 40:
            rw = c - 20;
            col = 1
        elif c >= 40 and c < 60:
            rw = c - 40;
            col = 2
        checks10[c].grid(row=rw, column=col)
    btn = tk.Button(root, text="Далее", bd=5, font=("Times New Roman", 12), width=10, height=2, command=get_entry_10)
    btn.grid(row=20, column=0, padx=20, pady=40)


def get_entry_10():
    for c in range(len(t10)):
        if var10[c].get() == False:
            t10[c] = ""
    for object in checks10:
        object.destroy()
    btn.destroy()
    prescription_4()


def prescription_4():
    global checks11, btn, final_btn, last_column
    last_column = 0
    checks11 = []
    for c in range(len(t11)):
        checks11.append(
            tk.Checkbutton(root, text=f"{t11[c][:60]}".replace("\n", ""), bg=maincolor, anchor="w", padx=10, width=52,
                           variable=var11[c],
                           onvalue=1, offvalue=0))
        if c < 20:
            rw = c;
            col = 0
        elif c >= 20 and c < 40:
            rw = c - 20;
            col = 1
        checks11[c].grid(row=rw, column=col)
    btn = tk.Button(root, text="Добавить", bd=5, font=("Times New Roman", 12), width=10, height=2, command=get_entry_11)
    btn.grid(row=20, column=0, padx=20, pady=40)
    final_btn = tk.Button(root, text="Готово", bd=5, font=("Times New Roman", 12), width=10, height=2, bg="#73a628",
                          command=ask_about_card)
    final_btn.grid(row=20, column=2, padx=20, pady=40)


def get_entry_11():
    global t12, last_column
    for c in range(len(t11)):
        if var11[c].get() and t11[c] != "":
            t12.append(t11[c])
            checks11[c].destroy()
            tk.Label(root, text=f"Добавлено: {t11[c][:26]}", bg=maincolor, anchor="w", padx=10, width=52).grid(
                row=last_column, column=2)
            t11[c] = ""
            last_column += 1


def ask_about_card():
    answer = mb.askyesno(title="Вопрос", message="Завести амбулаторную карту?")
    if answer:
        root.destroy()
        card_creation()
    else:
        root.destroy()


# create card
def card_creation():
    global root, ent_sex, ent_region, ent_district, ent_city, ent_street, ent_tel
    root = tk.Tk()
    root.title("Амбулаторная карта")
    root.geometry("1280x640+40+40")
    root.resizable(False, False)
    maincolor = "#A8DE56"
    root.config(bg=maincolor)

    lbl_sex = tk.Label(root, text="Пол:", bg=maincolor, font=("Times New Roman", 12))
    ent_sex = tk.Entry(root, font=("Times New Roman", 12), width=8)
    lbl_region = tk.Label(root, text="Область:", bg=maincolor, font=("Times New Roman", 12))
    ent_region = tk.Entry(root, font=("Times New Roman", 12), width=50)
    ent_region.insert(0, "Пензенская обл.")
    lbl_district = tk.Label(root, text="Район:", bg=maincolor, font=("Times New Roman", 12))
    ent_district = tk.Entry(root, font=("Times New Roman", 12), width=50)
    lbl_city = tk.Label(root, text="Населённый пункт:", bg=maincolor, font=("Times New Roman", 12))
    ent_city = tk.Entry(root, font=("Times New Roman", 12), width=50)
    ent_city.insert(0, "г. Пенза")
    lbl_street = tk.Label(root, text="Адрес (ул. Улица, дом-кв.):", bg=maincolor, font=("Times New Roman", 12))
    ent_street = tk.Entry(root, font=("Times New Roman", 12), width=50)
    lbl_tel = tk.Label(root, text="Телефон:", bg=maincolor, font=("Times New Roman", 12))
    ent_tel = tk.Entry(root, font=("Times New Roman", 12), width=50)
    btn = tk.Button(root, text="Готово", bd=5, width=10, height=2, font=("Times New Roman", 12), command=get_entry_card)

    lbl_sex.grid(row=0, column=0, padx=80, pady=5, sticky="w")
    ent_sex.grid(row=0, column=1, pady=5, sticky="w")
    lbl_region.grid(row=1, column=0, padx=80, pady=5, sticky="w")
    ent_region.grid(row=1, column=1, pady=5, sticky="w")
    lbl_district.grid(row=2, column=0, padx=80, pady=5, sticky="w")
    ent_district.grid(row=2, column=1, pady=5, sticky="w")
    lbl_city.grid(row=3, column=0, padx=80, pady=5, sticky="w")
    ent_city.grid(row=3, column=1, pady=5, sticky="w")
    lbl_street.grid(row=4, column=0, padx=80, pady=5, sticky="w")
    ent_street.grid(row=4, column=1, pady=5, sticky="w")
    lbl_tel.grid(row=5, column=0, padx=80, pady=5, sticky="w")
    ent_tel.grid(row=5, column=1, pady=5, sticky="w")
    ent_tel.grid(row=5, column=1, pady=5, sticky="w")
    btn.grid(row=6, column=0, padx=80, pady=5, sticky="w")


def get_entry_card():
    global sex, registration, add_card
    add_card = True
    sex = ent_sex.get().strip()
    if sex == "м" or sex == "м." or sex == "муж":
        sex = "муж."
    if sex == "ж" or sex == "ж." or sex == "жен":
        sex = "жен."
    region = ent_region.get().strip()
    if region != "":
        region = region + ", "
    district = ent_district.get().strip()
    if district != "":
        district = district + ", "
    city = ent_city.get().strip()
    if city != "":
        city = city + ", "
    street = ent_street.get().strip()
    if street != "":
        street = street + ", "
    tel = ent_tel.get().replace(" ", "")
    if len(tel) == 6:
        tel = f"тел.: (+7 8412) {tel[:2]}-{tel[2:4]}-{tel[4:6]}"
    elif len(tel) == 10:
        tel = f"тел.: +7 {tel[:3]} {tel[3:6]}-{tel[6:8]}-{tel[8:10]}"
    elif len(tel) == 11:
        tel = f"тел.: +7 {tel[1:4]} {tel[4:7]}-{tel[7:9]}-{tel[9:11]}"
    else:
        tel = f"тел.: {tel}"
    registration = region + district + city + street + tel
    root.destroy()


# main body
lbl_name = tk.Label(root, text="Пациент:", bg=maincolor, font=("Times New Roman", 12))
ent_name = tk.Entry(root, font=("Times New Roman", 12), width=50)
lbl_birthday = tk.Label(root, text="Дата рождения:", bg=maincolor, font=("Times New Roman", 12))
ent_birthday = tk.Entry(root, font=("Times New Roman", 12), width=50)
lbl_petition = tk.Label(root, text="Жалобы:", bg=maincolor, font=("Times New Roman", 12))
ent_petition = tk.Entry(root, font=("Times New Roman", 12), width=92)
lbl_head = tk.Label(root, text="Голова, см:", bg=maincolor, font=("Times New Roman", 12))
ent_head = tk.Entry(root, font=("Times New Roman", 12), width=8)
lbl_chest = tk.Label(root, text="Грудь, см:", bg=maincolor, font=("Times New Roman", 12))
ent_chest = tk.Entry(root, font=("Times New Roman", 12), width=8)
lbl_fonticulus = tk.Label(root, text="Большой родничок, мм:", bg=maincolor, font=("Times New Roman", 12))
ent_fonticulus = tk.Entry(root, font=("Times New Roman", 12), width=8)
lbl_allergy = tk.Label(root, text="Аллергические реакции:", bg=maincolor, font=("Times New Roman", 12))
ent_allergy = tk.Entry(root, font=("Times New Roman", 12), width=50)
btn = tk.Button(root, text="Далее", bd=5, width=10, height=2, font=("Times New Roman", 12), command=get_entry_start)

lbl_name.grid(row=0, column=0, padx=80, pady=5, sticky="w")
ent_name.grid(row=0, column=1, pady=5, sticky="w")
lbl_birthday.grid(row=1, column=0, padx=80, pady=5, sticky="w")
ent_birthday.grid(row=1, column=1, pady=5, sticky="w")
lbl_petition.grid(row=2, column=0, padx=80, pady=5, sticky="w")
ent_petition.grid(row=3, column=0, padx=80, pady=5, columnspan=2)
lbl_head.grid(row=4, column=0, padx=80, pady=5, sticky="w")
ent_head.grid(row=4, column=1, pady=5, sticky="w")
lbl_chest.grid(row=5, column=0, padx=80, pady=5, sticky="w")
ent_chest.grid(row=5, column=1, pady=5, sticky="w")
lbl_fonticulus.grid(row=6, column=0, padx=80, pady=5, sticky="w")
ent_fonticulus.grid(row=6, column=1, pady=5, sticky="w")
lbl_allergy.grid(row=7, column=0, padx=80, pady=5, sticky="w")
ent_allergy.grid(row=7, column=1, pady=5, sticky="w")
btn.grid(row=8, column=0, padx=80, pady=5, sticky="w")

root.mainloop()

# page config
document = Document()
section = document.sections[0]
section.page_height = Inches(11.69)
section.page_width = Inches(8.27)
section = document.sections[-1]
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)
section.left_margin = Inches(0.6)
section.right_margin = Inches(0.6)
paragraph_format = document.styles["Normal"].paragraph_format
paragraph_format.line_spacing = Pt(13)
style = document.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(10)

# date in dd.mm.yyyy
current_datetime = datetime.now()
if current_datetime.day > 9:
    day = f"{current_datetime.day}"
else:
    day = f"0{current_datetime.day}"
if current_datetime.month > 9:
    month = f"{current_datetime.month}"
else:
    month = f"0{current_datetime.month}"
year = f"{current_datetime.year}"
date = day + '.' + month + '.' + year

# decryption t2- with indexes(!)
if t2[5] != '' or t2[6] != '':
    for index in (7, 8, 9, 10, 11):
        t2[index] = t2[index].replace("\n", "")

# decryption t3
for t in range(last_hightonus_index, 1, -1):
    if t3[t] != '':
        t3[t] = t3[t].replace(",", ".")
        break

register = True
for t in range(1, hightonus_index):
    if t3[t] != '':
        register = False
if register:
    t3[hightonus_index] = t3[hightonus_index].capitalize()

# decryption t5
for t in range(len(t5)):
    if t5[t] == "Безусловные рефлексы: поисковый, сосательный, итд. ":
        t5[t] = (
                "\nБезусловные рефлексы: поисковый, сосательный, хоботковый, ладонно-рото-головной, хватательный, Моро\n"
                + "опоры, автоматической ходьбы, защитный, ползания, Галанта, Переса- симметричны. ")

# decryption t6
for t in range(len(t6)):
    if t6[t] == "Синдром мышечной дистонии: с кодом по МКБ. " or t6[
        t] == "Синдром мышечной дистонии: без кода по МКБ. ":
        t6[t] = "Синдром мышечной дистонии. "

# decryption t8
for t in range(len(t8)):
    if t8[t] == "консультация в Медико-генетическом научном центре\n":
        t8[
            t] = "консультация в Медико-генетическом научном центре им. академика Н.П. Бочкова (г. Москва, ул. Москворечье, 1)\n"
    if t8[t] == "консультации в НИИ педиатрии\n":
        t8[t] = (
                "консультация в НИИ педиатрии им. академика Ю.Е. Вельтищева ФГАОУ ВО РНИМУ им. Н.И. Пирогова\n"
                + "        Минздрава России (г. Москва, ул. Талдомская, 2)\n")

# decryption t9
for t in range(len(t9)):
    if t9[t] == "чередовать ванны: крапива, ромашка, череда\n":
        t9[t] = "чередовать ванны: крапива, ромашка, череда- 2 раза в неделю до № 10-15\n"
    if t9[t] == "успокаивающие ванны\n":
        t9[
            t] = "успокаивающие ванны: хмель, или мелисса, или лаванда, или ромашка, или душица- до № 15, затем сменить траву\n"
    if t9[t] == "ванны с крапивой\n":
        t9[t] = (
                "ванны с крапивой: 2 раза в неделю до № 10-15\n"
                + "        заварить 2 ст. л. сушёной крапивы на маленькую ванну, 4 ст. л. на большую ванну\n")
    if t9[t] == "ванны с морской солью\n":
        t9[t] = (
                "ванны с морской солью: 2 раза в неделю до № 10-15\n"
                + "        5 грамм морской соли на литр воды (75 г на маленькую ванну), температура воды до 37 °С:\n"
                + "        ребёнка искупать, промакнуть, через 20 минут соль смыть с тела\n")
    if t9[t] == "солёно-хвойные ванны\n":
        t9[t] = (
                "солёно-хвойные ванны: 2 раза в неделю до № 10-15\n"
                + "        5 грамм морской соли на литр воды (75 г на маленькую ванну) и 2 капли хвойного масла,\n"
                + "        температура воды до 37 °С: ребёнка искупать, промакнуть, через 20 минут соль смыть с тела\n")

# decryption t10
paraffin = (
        "        (при температуре окружающей среды до 25 °С, при температуре тела ребёнка до 37 °С):\n"
        + "        парафин (2 части: 400 г) и озокерит (1 часть: 200 г) растопить на водяной бане (осторожно! парафин горит!),\n"
        + "        вылить в форму, выстеленную медицинской клеёнкой, слоем в 1.5-2.0 см толщиной, остудить до 40 °С,\n"
        + "        наложить на 20 минут, затем снять (при беспокойстве ребёнка- снять раньше) и надеть тёплые носочки, варежки\n")

for t in range(len(t10)):
    if t10[t] == "парафин на кисти № 5+5\n" and t10[t + 1] == '':
        t10[t] = t10[t] + paraffin
    if t10[t] == "парафиновые сапожки № 10\n":
        t10[t] = t10[t] + paraffin

if t10[0] != "":
    for last_word_index in range(common_index, 0, -1):
        if t10[last_word_index] != "":
            t10[last_word_index] = t10[last_word_index].replace(", ", " № 10\n")
            break

# decryption t12
for t in range(len(t12)):
    if t12[t] == "мочегонный чай":
        t12[t] = (
                "мочегонный чай: 1 месяц\n"
                + "        лист земляники лесной, лист брусники, лист чёрной смородины, ромашку, плоды шиповника смешать в равных\n"
                + "        долях, заварить 1 чайную ложку смеси в 100 мл кипятка, пить по 1 чайной ложке чая 3 раза в день\n")
    elif t12[t] == "лист брусники":
        t12[t] = (
                "лист брусники: 1 месяц\n"
                + "        заварить 1 чайную ложку листа брусники в 100 мл кипятка, настоять, пить по 3 мл 2 раза в день 1 месяц\n")
    elif t12[t] == "пантогам (Hopantenic acid) 100 мг/мл в питьевом растворе":
        t12[t] = (t12[
                      t] + ":\n        7 дней: _______________________\n        21 день: ______________________\n        7 дней: _______________________\n")
    elif t12[t] == "пантогам или пантокальцин (Hopantenic acid) 250 мг":
        t12[t] = (t12[
                      t] + ":\n        7 дней: _______________________\n        21 день: ______________________\n        7 дней: _______________________\n")
    else:
        t12[t] = (t12[t] + ":   " + "_" * (95 - len(t12[t])) + "\n")

# title
document.add_heading(
    "ИП Павлова Ольга Игоревна\nгор. Пенза, ул. Пушкина, 15\nтел.: (+7 8412) 99-44-40 и +7 962 399-95-89\n"
    + "e-mail: rastuzdorovim@yandex.ru\nИНН: 583606020397   ОГРНИП: 315583600003199\n\n", level=3)
p = document.add_paragraph()
p.add_run("Дата:   ").bold = True
p.add_run(date + "\n")
p.add_run("Пациент:   ").bold = True
p.add_run(name + "\n")
p.add_run("Дата рождения:   ").bold = True
p.add_run(birthday + "\n\n\n")
p.add_run("Жалобы:   ").bold = True
if petition == "":
    p.add_run("Активных жалоб нет.\n\n")
p.add_run(petition + "\n\n")
p.add_run("Получил лечение:   __________________________________________").bold = True

p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
if head != "" or chest != "":
    if head != "":
        p.add_run("ОГ: ").bold = True
        p.add_run(head + " см")
    if chest != "":
        p.add_run("   ОГР: ").bold = True
        p.add_run(chest + " см")
    if fonticulus != "" and fonticulus.isdigit():
        p.add_run("   БР: ").bold = True
        p.add_run(fonticulus + "x" + fonticulus + " мм, не напряжён")
    elif fonticulus != "":
        p.add_run("   БР: ").bold = True
        p.add_run(fonticulus)
if allergy == "":
    p.add_run("\nНаличие аллергических реакций отрицают")
p.add_run("\n" + allergy)

# final text from dictionary
p = document.add_paragraph()
p.add_run("Неврологический статус:   ").bold = True
p.add_run("В сознании. Общемозговой симптоматики нет. Рвоты, судорог при осмотре нет. ")
for text in t1:
    p.add_run(text)
p.add_run("\nЧН:   ").bold = True
for t in [t2, t3, t4, t5]:
    for text in t:
        p.add_run(text)
    p.add_run("\n")
p.add_run("\n\n")
p.add_run("Диагноз поставлен на основании жалоб, объективных данных и лабораторных обследований.\n\n")
if t6[0] == "Здоров. ":
    p.add_run("В неврологическом статусе: без очаговых знаков.\n")
p.add_run("Диагноз:   ").bold = True
for text in t6:
    p.add_run(text)
p.add_run("\nДиагноз по МКБ-10:   ").bold = True
for text in t7:
    p.add_run(text)

document.add_page_break()
p = document.add_paragraph()
p.add_run("\nРекомендовано обследование:   \n").bold = True
for text in t8:
    p.add_run(text)
p.add_run("\n")
p.add_run("Рекомендовано лечение:   \n").bold = True
for t in [t9, t12, t10]:
    for text in t:
        p.add_run(text)
p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run("Врач-невролог:   ").bold = True
p.add_run("Павлова Ольга Игоревна")
p = document.add_paragraph()
p.add_run("\n\nЯвка на повторный осмотр:   ____________________________\n\nЭпидемиологический анамнез:   ").bold = True
p.add_run(
    "Нахождение в очаге инфекции мама отрицает. Наличие контактов с больными\n"
    + "ОРВИ с инфекцией, вызванной новым коронавирусом SARS-CoV2, за последние 14 дней отрицает. Наличие тесных\n"
    + "контактов с лицами, находящимися под наблюдением по инфекции, вызванной новым коронавирусом SARS-CoV2,\n"
    + "которые в последующем заболели, отрицает. Наличие тесных контактов с лицами, у которых лабораторно\n"
    + "подтверждён  COVID-19, отрицает. За пределы Пензенской области и РФ в течение последнего месяца не выезжали.\n"
    + "В контакте с лицами, вернувшимися из других регионов и стран, не были.\n"
    + "TBC, ВИЧ, сифилис в семье и у родственников отрицает.\n"
    + "Дисфункций кишечника за последние 3 недели не было.\n\n"
    + "Выписка выдана на руки.\n"
    + "С планом обследования и лечения согласен (-на), о последствиях предупреждён (-на):   __________________________")
p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run("пациент (законный представитель)\n")
p.add_run(date)

# save and open file
sname = name.split()
savename = f"{sname[0]} {sname[1]} {date[:6]}{date[8:]}"
document.save(f"C:\\Users\\Кабинет 3\\Desktop\\Расту здоровым\\Архив\\{savename}.docx")
os.startfile(f"C:\\Users\\Кабинет 3\\Desktop\\Расту здоровым\\Архив\\{savename}.docx")

# print card
if add_card:
    f = open("card_number.txt", "r+")
    card_number = f.read()
    f.seek(0)
    f.write(f"{int(card_number) + 1}")
    f.close()

    card = Document("outpatient_card.docx")
    style = card.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)

    par = card.paragraphs
    par[15].add_run(card_number).bold = True
    par[19].add_run(date)
    par[20].add_run(name)
    par[21].add_run(sex)
    par[22].add_run(birthday)
    par[23].add_run(registration)
    card.save(
        f"C:\\Users\\Кабинет 3\\Desktop\\Расту здоровым\\Амбулаторные карты\\Карта № {card_number} {sname[0]} {sname[1]}.docx")
    os.startfile(
        f"C:\\Users\\Кабинет 3\\Desktop\\Расту здоровым\\Амбулаторные карты\\Карта № {card_number} {sname[0]} {sname[1]}.docx")
